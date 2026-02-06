
import sys
import os
import io

# Force UTF-8 output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def install_and_import():
    try:
        import docx
        return docx
    except ImportError:
        print("正在安裝 python-docx...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
        return docx

def read_docx(file_path):
    docx = install_and_import()
    from docx import Document
    
    if not os.path.exists(file_path):
        print(f"錯誤: 找不到檔案 {file_path}")
        return

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"無法讀取檔案 (可能是格式錯誤或加密): {e}")
        return

    output = []
    
    # 1. 讀取段落
    output.append(f"=== 檔案: {os.path.basename(file_path)} ===")
    output.append("--- 段落內容 ---")
    for para in doc.paragraphs:
        if para.text.strip():
            output.append(para.text)
            
    # 2. 讀取表格
    output.append("\n--- 表格內容 ---")
    for i, table in enumerate(doc.tables):
        output.append(f"[表格 {i+1}]")
        for row in table.rows:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            output.append(" | ".join(row_data))
        output.append("") # 表格空行

    result = "\n".join(output)
    print(result)
    return result

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python read_docx.py <path_to_docx>")
        print("例如: python read_docx.py my_document.docx")
    else:
        read_docx(sys.argv[1])
